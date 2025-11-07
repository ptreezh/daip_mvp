"""
DOCX to Markdown converter implementation.
"""
import asyncio
from pathlib import Path
from typing import Any
import aiofiles
import logging

from daip_live.doc.converter.base_converter import BaseConverter
from daip_live.doc.models.document_models import DocumentConversionResult


class DocxToMarkdownConverter(BaseConverter):
    """Convert DOCX files to Markdown format."""
    
    def __init__(self):
        super().__init__()
        self.supported_formats = {".docx", ".doc"}
        self.name = "DOCX To Markdown Converter"
        self.logger = logging.getLogger(__name__)
    
    async def convert(self, source_path: Path, target_path: Path) -> DocumentConversionResult:
        """Convert DOCX to markdown format."""
        start_time = asyncio.get_event_loop().time()
        
        result = DocumentConversionResult(
            source_format="docx",
            target_format="markdown",
            source_path=str(source_path),
            target_path=str(target_path),
            success=False
        )
        
        try:
            # Validate input
            if not await self.validate_input(source_path):
                result.error_message = f"Source file does not exist: {source_path}"
                result.conversion_time = asyncio.get_event_loop().time() - start_time
                return result
            
            # Ensure target path has .md extension
            if target_path.suffix.lower() not in ['.md', '.markdown']:
                target_path = target_path.with_suffix('.md')
            
            # Import python-docx only when needed
            from docx import Document
            
            # Read DOCX content
            doc = Document(str(source_path))
            
            # Convert to markdown
            markdown_content = self._process_docx_content(doc)
            
            # Write markdown content
            async with aiofiles.open(target_path, 'w', encoding='utf-8') as f:
                await f.write(markdown_content)
            
            # Validate output
            success = await self.validate_output(target_path)
            
            # Calculate size
            size = target_path.stat().st_size if target_path.exists() else 0
            
            end_time = asyncio.get_event_loop().time()
            
            result.success = success
            result.target_path = str(target_path)
            result.converted_size = size
            result.conversion_time = end_time - start_time
            
        except Exception as e:
            end_time = asyncio.get_event_loop().time()
            result.error_message = str(e)
            result.conversion_time = end_time - start_time
            self.logger.error(f"Error converting docx to markdown: {e}")
        
        return result
    
    def _process_docx_content(self, doc) -> str:
        """Process DOCX document and extract as markdown."""
        markdown_elements = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            # Identify heading levels based on style
            if paragraph.style and paragraph.style.name.startswith('Heading'):
                try:
                    # Extract heading level from style name
                    level_str = ''.join(filter(str.isdigit, paragraph.style.name))
                    if level_str:
                        level = max(1, min(6, int(level_str)))  # Limit to H1-H6
                        markdown_elements.append(f"{'#' * level} {text}")
                    else:
                        markdown_elements.append(text)  # Regular paragraph
                except:
                    markdown_elements.append(text)  # Fallback to regular paragraph
            elif paragraph.style and paragraph.style.name in ['List Bullet', 'List Number']:
                # List items
                if 'Bullet' in paragraph.style.name:
                    markdown_elements.append(f"- {text}")
                elif 'Number' in paragraph.style.name:
                    markdown_elements.append(f"1. {text}")
            elif paragraph.style and 'Quote' in paragraph.style.name:
                # Block quotes
                markdown_elements.append(f"> {text}")
            else:
                # Regular paragraphs
                markdown_elements.append(text)
        
        # Process tables if any
        for table in doc.tables:
            if len(table.rows) > 0 and table.cell(0, 0).text.strip():  # If table has content
                markdown_elements.append("\n")  # Add spacing before table
                # Create markdown table representation
                for i, row in enumerate(table.rows):
                    cells = [cell.text.strip() for cell in row.cells]
                    row_str = "| " + " | ".join(cells) + " |"
                    markdown_elements.append(row_str)
                    
                    # Add separator after first row (headers)
                    if i == 0:
                        sep_row = "| " + " | ".join(['---'] * len(cells)) + " |"
                        markdown_elements.append(sep_row)
        
        return '\n\n'.join(markdown_elements)