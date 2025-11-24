"""
Markdown to DOCX converter implementation.
"""
import asyncio
from pathlib import Path
from typing import Any
import aiofiles
import logging

from docx import Document
from docx.shared import Inches
from markdown import markdown
from bs4 import BeautifulSoup

from daip_live.doc.converter.base_converter import BaseConverter
from daip_live.doc.models.document_models import DocumentConversionResult


class MarkdownToDocxConverter(BaseConverter):
    """Convert markdown files to DOCX format."""
    
    def __init__(self):
        super().__init__()
        self.supported_formats = {".md", ".markdown"}
        self.name = "Markdown To DOCX Converter"
        self.logger = logging.getLogger(__name__)
    
    async def convert(self, source_path: Path, target_path: Path) -> DocumentConversionResult:
        """Convert markdown to DOCX format."""
        start_time = asyncio.get_event_loop().time()
        
        result = DocumentConversionResult(
            source_format="markdown",
            target_format="docx",
            source_path=str(source_path),
            target_path=str(target_path),
            success=False
        )
        
        try:
            # Validate input
            if not await self.validate_input(source_path):
                result.error_message = f"Source file does not exist: {source_path}"
                result.conversion_time = asyncio.get_event_loop().time() - start_time
                return result
            
            # Ensure target path has .docx extension
            if target_path.suffix.lower() != '.docx':
                target_path = target_path.with_suffix('.docx')
            
            # Read markdown content
            async with aiofiles.open(source_path, 'r', encoding='utf-8') as f:
                content = await f.read()
            
            # Create new document
            doc = Document()
            
            # Process content using enhanced markdown parsing
            self._process_markdown_content(content, doc)
            
            # Save document
            doc.save(str(target_path))
            
            # Validate output
            success = await self.validate_output(target_path)
            
            # Calculate size
            size = target_path.stat().st_size if target_path.exists() else 0
            
            end_time = asyncio.get_event_loop().time()
            
            result.success = success
            result.target_path = str(target_path)
            result.converted_size = size
            result.conversion_time = end_time - start_time
            
        except Exception as e:
            end_time = asyncio.get_event_loop().time()
            result.error_message = str(e)
            result.conversion_time = end_time - start_time
            self.logger.error(f"Error converting markdown to docx: {e}")
        
        return result
    
    def _process_markdown_content(self, content: str, doc) -> None:
        """Process markdown content and add to DOCX document."""
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                # Empty line - add spacing
                doc.add_paragraph("")
                continue
            
            # Process different markdown elements
            if line.startswith('# '):
                # Heading level 1
                doc.add_heading(line[2:], level=0)
            elif line.startswith('## '):
                # Heading level 2
                doc.add_heading(line[3:], level=1)
            elif line.startswith('### '):
                # Heading level 3
                doc.add_heading(line[4:], level=2)
            elif line.startswith('#### '):
                # Heading level 4
                doc.add_heading(line[5:], level=3)
            elif line.startswith('- ') or line.startswith('* ') or line.startswith('+ '):
                # Bullet point
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(line[2:])
            elif line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                # Numbered list (basic implementation)
                p = doc.add_paragraph(style='List Number')
                # Find the content after the number
                # Find first dot to extract number and content
                dot_idx = line.find('.')
                if dot_idx != -1:
                    content_start = dot_idx + 2  # Skip "N. "
                    if content_start < len(line):
                        p.add_run(line[content_start:])
            elif line.startswith('> '):
                # Block quote
                p = doc.add_paragraph(line[2:])
                p.style = 'Intense Quote'  # Use proper style name if available
            elif line.startswith('**') and line.endswith('**') and len(line) > 4:
                # Bold text
                p = doc.add_paragraph()
                run = p.add_run(line[2:-2])
                run.bold = True
            elif line.startswith('*') and line.endswith('*') and len(line) > 2 and not (line.startswith('**') and line.endswith('**')):
                # Italic text (not including bold)
                p = doc.add_paragraph()
                run = p.add_run(line[1:-1])
                run.italic = True
            elif line.startswith('```'):
                # Code block
                p = doc.add_paragraph(line[3:])  # Just add as regular text for now
                # In the future, could apply a monospace font
            elif '|' in line and '+' in line and '--' in line:
                # Table detection (simplified)
                p = doc.add_paragraph(line)
            else:
                # Regular paragraph
                p = doc.add_paragraph()
                p.add_run(line)


class DocxToMarkdownConverter(BaseConverter):
    """Convert DOCX files to Markdown format."""

    def __init__(self):
        super().__init__()
        self.supported_formats = {".docx", ".doc"}
        self.name = "DOCX To Markdown Converter"
        self.logger = logging.getLogger(__name__)
    
    async def convert(self, source_path: Path, target_path: Path) -> DocumentConversionResult:
        """Convert DOCX to markdown format."""
        start_time = asyncio.get_event_loop().time()
        
        result = DocumentConversionResult(
            source_format="docx",
            target_format="markdown",
            source_path=str(source_path),
            target_path=str(target_path),
            success=False
        )
        
        try:
            # Validate input
            if not await self.validate_input(source_path):
                result.error_message = f"Source file does not exist: {source_path}"
                result.conversion_time = asyncio.get_event_loop().time() - start_time
                return result
            
            # Ensure target path has .md extension
            if target_path.suffix.lower() not in ['.md', '.markdown']:
                target_path = target_path.with_suffix('.md')
            
            # Import python-docx only when needed
            from docx import Document
            
            # Read DOCX content
            doc = Document(str(source_path))
            
            # Convert to markdown
            markdown_content = await self._process_docx_content(doc)
            
            # Write markdown content
            async with aiofiles.open(target_path, 'w', encoding='utf-8') as f:
                await f.write(markdown_content)
            
            # Validate output
            success = await self.validate_output(target_path)
            
            # Calculate size
            size = target_path.stat().st_size if target_path.exists() else 0
            
            end_time = asyncio.get_event_loop().time()
            
            result.success = success
            result.target_path = str(target_path)
            result.converted_size = size
            result.conversion_time = end_time - start_time
            
        except Exception as e:
            end_time = asyncio.get_event_loop().time()
            result.error_message = str(e)
            result.conversion_time = end_time - start_time
            self.logger.error(f"Error converting docx to markdown: {e}")
        
        return result
    
    async def _process_docx_content(self, doc) -> str:
        """Process DOCX document and extract as markdown."""
        markdown_elements = []
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            
            # Identify heading levels based on style
            if paragraph.style.name.startswith('Heading'):
                level = 1
                try:
                    # Extract heading level from style name
                    level_str = paragraph.style.name.replace('Heading', '').strip()
                    if level_str.isdigit():
                        level = int(level_str)
                    else:
                        level = 1
                except:
                    level = 1
                
                markdown_elements.append(f"{'#' * level} {text}")
            elif paragraph.style.name in ['List Bullet', 'List Number']:
                # List items
                if paragraph.style.name == 'List Bullet':
                    markdown_elements.append(f"- {text}")
                else:
                    markdown_elements.append(f"1. {text}")
            elif paragraph.style.name == 'Intense Quote':
                # Block quotes
                markdown_elements.append(f"> {text}")
            else:
                # Regular paragraphs
                markdown_elements.append(text)
        
        # Process tables if any
        for table in doc.tables:
            if table.cell(0, 0).text.strip():  # If table has content
                markdown_elements.append("\n| Column 1 | Column 2 |")  # Template
                markdown_elements.append("|----------|----------|")
                # We'll add a simplified representation
        
        return '\n\n'.join(markdown_elements)